"""
AI Resume Analyzer - Flask Backend
===================================
Features:
- Resume parsing (PDF & DOCX)
- ATS Scoring
- Skill extraction
- TF-IDF + Cosine Similarity
- Job matching
- Skill gap analysis
- Learning roadmap generation
- PDF report export
"""

from flask import Flask, render_template, request, jsonify, session, send_file
from flask_cors import CORS
import os, re, io, json, uuid, math
from datetime import datetime
from collections import Counter
from flask_session import Session
from flask_mysqldb import MySQL

# ── PDF / DOCX parsing ──────────────────────────────────────────────────────
try:
    import pdfplumber
    PDF_OK = True
except ImportError:
    PDF_OK = False

try:
    from docx import Document as DocxDocument
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# ── NLP / ML ─────────────────────────────────────────────────────────────────
try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    import numpy as np
    ML_OK = True
except ImportError:
    ML_OK = False

# ── PDF Generation ────────────────────────────────────────────────────────────
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, HRFlowable, KeepTogether)
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False

app = Flask(__name__)
app.secret_key = "resumeiq-secret-2024"
app.config['SESSION_TYPE'] = 'filesystem'
app.config['SESSION_PERMANENT'] = False
app.config['SESSION_FILE_DIR'] = './flask_session'
Session(app)
CORS(app, supports_credentials=True)
app.config['SESSION_COOKIE_SAMESITE'] = "None"
app.config['SESSION_COOKIE_SECURE'] = True

app.config['MYSQL_HOST'] = 'hopper.proxy.rlwy.net'
app.config['MYSQL_USER'] = 'root'
app.config['MYSQL_PASSWORD'] = 'HadHlabRNbLxrJPqgqgkqukbmKuAQzVy'
app.config['MYSQL_DB'] = 'railway'
app.config['MYSQL_PORT'] = 55341
app.config['MYSQL_CHARSET'] = 'utf8mb4'
app.config['MYSQL_COLLATION'] = 'utf8mb4_unicode_ci'
mysql = MySQL(app)

UPLOAD_FOLDER = os.path.join(os.getcwd(), "uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# ═══════════════════════════════════════════════════════════════════════════════
#  LARGE DATASET
# ═══════════════════════════════════════════════════════════════════════════════

JOB_DATABASE = {
    # ── Data Science & AI ──────────────────────────────────────────────────
    "Data Scientist": {
        "skills": ["python","machine learning","deep learning","statistics","pandas","numpy","scikit-learn",
                   "sql","data visualization","tensorflow","keras","pytorch","r","feature engineering",
                   "hypothesis testing","a/b testing","tableau","power bi","spark","hadoop","jupyter",
                   "matplotlib","seaborn","nlp","computer vision","time series","regression","classification",
                   "clustering","random forest","xgboost","lightgbm","model deployment","mlflow","airflow"],
        "avg_salary": "₹12–25 LPA",
        "demand": "Very High",
        "description": "Analyze complex datasets to derive business insights using ML/AI techniques.",
        "companies": ["Google","Amazon","Flipkart","Swiggy","Zomato","Infosys","TCS","Wipro","Myntra"]
    },
    "Machine Learning Engineer": {
        "skills": ["python","machine learning","deep learning","tensorflow","pytorch","keras","scikit-learn",
                   "docker","kubernetes","git","rest api","mlflow","feature engineering","model deployment",
                   "ci/cd","cloud computing","aws","gcp","azure","spark","kafka","model optimization",
                   "onnx","triton","cuda","distributed training","transformers","huggingface","fastapi",
                   "flask","sql","nosql","redis","model monitoring","a/b testing","data pipelines"],
        "avg_salary": "₹15–35 LPA",
        "demand": "Very High",
        "description": "Build, deploy, and maintain production-grade ML systems at scale.",
        "companies": ["Microsoft","Google","Amazon","Meta","Apple","Uber","Airbnb","Razorpay","Paytm"]
    },
    "AI Research Scientist": {
        "skills": ["python","deep learning","pytorch","tensorflow","mathematics","linear algebra",
                   "calculus","probability","statistics","nlp","computer vision","reinforcement learning",
                   "transformers","huggingface","research paper reading","latex","c++","cuda",
                   "distributed systems","optimization","signal processing","graph neural networks",
                   "generative ai","diffusion models","llm fine-tuning","rlhf","arxiv","publications"],
        "avg_salary": "₹20–50 LPA",
        "demand": "High",
        "description": "Conduct cutting-edge AI research and publish novel algorithms.",
        "companies": ["Google DeepMind","OpenAI","Meta AI","Microsoft Research","IISc","IIT Labs","TCS Research"]
    },
    "NLP Engineer": {
        "skills": ["python","nlp","transformers","huggingface","bert","gpt","spacy","nltk","text classification",
                   "named entity recognition","sentiment analysis","machine translation","question answering",
                   "information extraction","pytorch","tensorflow","fastapi","docker","elasticsearch",
                   "vector databases","langchain","llm","prompt engineering","rag","fine-tuning","tokenization",
                   "word embeddings","word2vec","glove","fasttext","regex","sql","mongodb"],
        "avg_salary": "₹14–30 LPA",
        "demand": "Very High",
        "description": "Build intelligent systems that understand and generate human language.",
        "companies": ["Sarvam AI","Krutrim","Google","Amazon","Freshworks","Uniphore","Observe.AI"]
    },
    "Computer Vision Engineer": {
        "skills": ["python","opencv","deep learning","pytorch","tensorflow","image segmentation","object detection",
                   "cnn","yolo","resnet","vgg","efficientnet","data augmentation","image preprocessing",
                   "c++","cuda","tensorrt","onnx","docker","git","rest api","3d vision","slam",
                   "lidar processing","point clouds","gan","stable diffusion","image generation","video analytics"],
        "avg_salary": "₹12–28 LPA",
        "demand": "High",
        "description": "Develop systems that extract meaning from visual data.",
        "companies": ["Ola","Nuro","Waymo","Intel","Qualcomm","SigTuple","Tricog Health"]
    },
    "Data Analyst": {
        "skills": ["sql","excel","python","tableau","power bi","statistics","data visualization",
                   "r","pandas","numpy","business intelligence","reporting","etl","google analytics",
                   "looker","data storytelling","pivot tables","vlookup","hypothesis testing","dashboards",
                   "google sheets","data cleaning","exploratory data analysis","mongodb","mysql","postgresql"],
        "avg_salary": "₹5–15 LPA",
        "demand": "Very High",
        "description": "Transform raw data into actionable business insights through analysis and visualization.",
        "companies": ["Deloitte","EY","KPMG","PwC","Accenture","Cognizant","Capgemini","MakeMyTrip"]
    },
    "Data Engineer": {
        "skills": ["python","sql","spark","hadoop","kafka","airflow","dbt","aws","gcp","azure",
                   "data warehousing","etl","snowflake","bigquery","redshift","delta lake","databricks",
                   "docker","kubernetes","git","scala","java","nosql","cassandra","hbase","hive",
                   "data modeling","schema design","data quality","orchestration","terraform","ci/cd"],
        "avg_salary": "₹10–28 LPA",
        "demand": "Very High",
        "description": "Design and build data pipelines and infrastructure at scale.",
        "companies": ["Zepto","Blinkit","PhonePe","CRED","Groww","Angel One","Delhivery"]
    },
    "Business Analyst": {
        "skills": ["sql","excel","tableau","power bi","requirements gathering","process modeling","uml",
                   "stakeholder management","jira","confluence","agile","scrum","data analysis",
                   "business intelligence","presentations","documentation","python","r","statistics",
                   "product management","wireframing","user stories","brd","frd","gap analysis","erp","sap"],
        "avg_salary": "₹6–18 LPA",
        "demand": "High",
        "description": "Bridge the gap between business needs and technical solutions.",
        "companies": ["Accenture","Infosys","Wipro","TCS","HDFC Bank","ICICI Bank","Reliance","Tata Group"]
    },
    "MLOps Engineer": {
        "skills": ["python","mlflow","kubeflow","airflow","docker","kubernetes","ci/cd","git","terraform",
                   "aws","gcp","azure","model monitoring","feature stores","model registry","jenkins",
                   "github actions","prometheus","grafana","elk stack","data versioning","dvc","evidently",
                   "seldon","bentoml","fastapi","spark","kafka","model deployment","a/b testing","shadow mode"],
        "avg_salary": "₹14–32 LPA",
        "demand": "Very High",
        "description": "Operationalize ML models with robust deployment and monitoring infrastructure.",
        "companies": ["Flipkart","Amazon","Zomato","Swiggy","Meesho","InMobi","Freshworks"]
    },
    # ── Software Engineering ──────────────────────────────────────────────
    "Full Stack Developer": {
        "skills": ["javascript","typescript","react","nodejs","python","html","css","sql","mongodb",
                   "rest api","graphql","git","docker","aws","redis","postgresql","mysql","express",
                   "nextjs","vuejs","angular","webpack","jest","ci/cd","agile","linux","nginx","elasticsearch"],
        "avg_salary": "₹7–22 LPA",
        "demand": "Very High",
        "description": "Build complete web applications from front-end to back-end and databases.",
        "companies": ["Startups","Razorpay","Zerodha","Notion","Atlassian","ThoughtWorks","Gojek"]
    },
    "Backend Developer": {
        "skills": ["python","java","golang","nodejs","sql","postgresql","mysql","mongodb","redis",
                   "rest api","microservices","docker","kubernetes","aws","gcp","kafka","rabbitmq",
                   "git","system design","ci/cd","linux","nginx","grpc","graphql","elasticsearch",
                   "caching","performance optimization","security","authentication","oauth"],
        "avg_salary": "₹8–25 LPA",
        "demand": "Very High",
        "description": "Design scalable server-side applications and APIs.",
        "companies": ["Razorpay","CRED","Zerodha","PhonePe","Dunzo","Urban Company","BrowserStack"]
    },
    "DevOps Engineer": {
        "skills": ["docker","kubernetes","jenkins","git","ansible","terraform","aws","gcp","azure",
                   "linux","bash","python","ci/cd","monitoring","prometheus","grafana","elk stack",
                   "helm","istio","service mesh","security","networking","nginx","load balancing",
                   "cloud architecture","cost optimization","site reliability","incident management"],
        "avg_salary": "₹10–28 LPA",
        "demand": "High",
        "description": "Bridge development and operations through automation and infrastructure.",
        "companies": ["Infosys","TCS","HCL","Tech Mahindra","Mindtree","Mphasis","Persistent"]
    },
    "Cloud Architect": {
        "skills": ["aws","gcp","azure","terraform","kubernetes","docker","microservices","security",
                   "networking","serverless","cost optimization","cloud migration","iaas","paas","saas",
                   "ci/cd","devops","linux","python","system design","high availability","disaster recovery",
                   "cloud native","multi-cloud","compliance","iam","vpc","load balancing","cdn"],
        "avg_salary": "₹18–45 LPA",
        "demand": "High",
        "description": "Design enterprise-scale cloud infrastructure and architecture.",
        "companies": ["Accenture","Deloitte","Amazon","Microsoft","Google","IBM","Capgemini"]
    },
    # ── Cybersecurity ─────────────────────────────────────────────────────
    "Cybersecurity Analyst": {
        "skills": ["network security","ethical hacking","penetration testing","siem","soc","incident response",
                   "vulnerability assessment","firewalls","ids/ips","python","linux","wireshark","metasploit",
                   "owasp","cryptography","compliance","iso 27001","gdpr","threat intelligence","forensics",
                   "endpoint security","cloud security","zero trust","identity management","scripting"],
        "avg_salary": "₹8–22 LPA",
        "demand": "Very High",
        "description": "Protect organizations from cyber threats and security breaches.",
        "companies": ["Wipro","HCL","Quick Heal","Lucideus","Sequretek","CERT-In","Banks & BFSI firms"]
    },
    # ── Product & Design ─────────────────────────────────────────────────
    "Product Manager": {
        "skills": ["product strategy","roadmap","agile","scrum","jira","user research","data analysis",
                   "sql","a/b testing","wireframing","figma","stakeholder management","go-to-market",
                   "metrics","okrs","user stories","competitive analysis","customer development",
                   "presentation","excel","python","product analytics","mixpanel","amplitude","growth"],
        "avg_salary": "₹15–40 LPA",
        "demand": "High",
        "description": "Define product vision and lead cross-functional teams to deliver value.",
        "companies": ["Meesho","Nykaa","CRED","Groww","BrowserStack","Freshworks","Zoho"]
    },
    "UI/UX Designer": {
        "skills": ["figma","sketch","adobe xd","user research","wireframing","prototyping","usability testing",
                   "design systems","html","css","interaction design","information architecture","typography",
                   "color theory","accessibility","responsive design","motion design","user interviews",
                   "a/b testing","heuristic evaluation","persona creation","journey mapping","adobe illustrator"],
        "avg_salary": "₹6–18 LPA",
        "demand": "High",
        "description": "Create intuitive and beautiful digital experiences for users.",
        "companies": ["Swiggy","Zomato","Dunzo","PhonePe","Ola","Byju's","Unacademy"]
    },
}

# ── ATS Keywords by category ────────────────────────────────────────────────
ATS_SECTIONS = {
    "Contact Information": ["email","phone","linkedin","github","portfolio","address","location"],
    "Education": ["bachelor","master","phd","b.tech","m.tech","degree","university","college","gpa","cgpa","graduation"],
    "Experience": ["experience","worked","internship","project","company","organization","team","led","managed","developed","built","designed","implemented","achieved","increased","reduced","improved"],
    "Technical Skills": ["python","java","javascript","sql","machine learning","deep learning","tensorflow","pytorch","docker","kubernetes","aws","gcp","azure","react","nodejs","git"],
    "Soft Skills": ["communication","leadership","teamwork","problem solving","critical thinking","adaptability","time management","collaboration"],
    "Certifications": ["certified","certification","aws certified","google certified","coursera","udemy","nptel","microsoft certified"],
    "Achievements": ["award","scholarship","winner","rank","topper","published","patent","hackathon","competition"],
    "Action Verbs": ["developed","designed","implemented","led","managed","created","built","deployed","optimized","collaborated","researched","analyzed","improved","achieved"]
}

# ── Learning Resources Database ──────────────────────────────────────────────
LEARNING_RESOURCES = {
    "python": {
        "courses": [
            {"name": "Python for Everybody – University of Michigan", "platform": "Coursera", "url": "https://coursera.org/specializations/python", "duration": "8 months", "level": "Beginner"},
            {"name": "Complete Python Bootcamp", "platform": "Udemy", "url": "https://udemy.com/course/complete-python-bootcamp", "duration": "22 hrs", "level": "Beginner"},
            {"name": "Python Data Science Handbook", "platform": "GitHub (Free)", "url": "https://jakevdp.github.io/PythonDataScienceHandbook", "duration": "Self-paced", "level": "Intermediate"},
        ],
        "projects": [
            "Build a web scraper with BeautifulSoup and store data in SQLite",
            "Create a REST API with FastAPI and PostgreSQL",
            "Develop a CLI tool for data processing with Pandas",
            "Build an automation bot using Python and Selenium",
        ],
        "certifications": [
            {"name": "PCEP – Python Certified Entry-Level", "org": "Python Institute", "url": "https://pythoninstitute.org"},
            {"name": "Python 3 Programming Specialization", "org": "Coursera (UMich)", "url": "https://coursera.org"},
        ],
        "timeline": "6–8 weeks for fundamentals, 3–6 months for proficiency"
    },
    "machine learning": {
        "courses": [
            {"name": "Machine Learning Specialization", "platform": "Coursera (Andrew Ng)", "url": "https://coursera.org/specializations/machine-learning-introduction", "duration": "3 months", "level": "Beginner-Intermediate"},
            {"name": "Machine Learning A–Z", "platform": "Udemy", "url": "https://udemy.com/course/machinelearning", "duration": "44 hrs", "level": "Intermediate"},
            {"name": "Hands-On ML with Scikit-Learn & TensorFlow", "platform": "O'Reilly (Book)", "url": "https://oreilly.com", "duration": "Self-paced", "level": "Intermediate"},
        ],
        "projects": [
            "House price prediction using regression (Kaggle Boston dataset)",
            "Customer churn prediction with classification algorithms",
            "Recommendation system using collaborative filtering",
            "Sentiment analysis on IMDB movie reviews",
            "Credit card fraud detection with imbalanced data",
        ],
        "certifications": [
            {"name": "TensorFlow Developer Certificate", "org": "Google", "url": "https://tensorflow.org/certificate"},
            {"name": "AWS Certified ML – Specialty", "org": "Amazon Web Services", "url": "https://aws.amazon.com/certification"},
            {"name": "Professional ML Engineer", "org": "Google Cloud", "url": "https://cloud.google.com/certification"},
        ],
        "timeline": "3–6 months of dedicated study"
    },
    "deep learning": {
        "courses": [
            {"name": "Deep Learning Specialization (5 courses)", "platform": "Coursera (deeplearning.ai)", "url": "https://coursera.org/specializations/deep-learning", "duration": "5 months", "level": "Intermediate"},
            {"name": "fast.ai Practical Deep Learning", "platform": "fast.ai (Free)", "url": "https://course.fast.ai", "duration": "7 weeks", "level": "Intermediate"},
            {"name": "NYU Deep Learning Course", "platform": "YouTube (Free)", "url": "https://atcold.github.io/pytorch-Deep-Learning", "duration": "Self-paced", "level": "Advanced"},
        ],
        "projects": [
            "Image classification using CNN on CIFAR-10",
            "Object detection with YOLO on custom dataset",
            "Text generation using LSTM / Transformer",
            "GAN-based image synthesis (face generation)",
            "Transfer learning for medical image diagnosis",
        ],
        "certifications": [
            {"name": "Deep Learning Specialization Certificate", "org": "deeplearning.ai / Coursera", "url": "https://coursera.org"},
            {"name": "TensorFlow Developer Certificate", "org": "Google", "url": "https://tensorflow.org/certificate"},
        ],
        "timeline": "4–8 months after knowing ML basics"
    },
    "sql": {
        "courses": [
            {"name": "SQL for Data Science", "platform": "Coursera (UC Davis)", "url": "https://coursera.org/learn/sql-for-data-science", "duration": "4 weeks", "level": "Beginner"},
            {"name": "The Complete SQL Bootcamp", "platform": "Udemy", "url": "https://udemy.com/course/the-complete-sql-bootcamp", "duration": "9 hrs", "level": "Beginner"},
            {"name": "Mode Analytics SQL Tutorial", "platform": "Mode (Free)", "url": "https://mode.com/sql-tutorial", "duration": "Self-paced", "level": "Intermediate"},
        ],
        "projects": [
            "Analyze sales data using window functions and CTEs",
            "Build a student database management system",
            "Perform e-commerce analytics with complex joins",
            "Create a real-time dashboard from a PostgreSQL database",
        ],
        "certifications": [
            {"name": "Oracle Database SQL Certified Associate", "org": "Oracle", "url": "https://education.oracle.com"},
            {"name": "Microsoft DP-900: Azure Data Fundamentals", "org": "Microsoft", "url": "https://learn.microsoft.com"},
        ],
        "timeline": "4–6 weeks for intermediate proficiency"
    },
    "docker": {
        "courses": [
            {"name": "Docker & Kubernetes: The Practical Guide", "platform": "Udemy", "url": "https://udemy.com/course/docker-kubernetes-the-practical-guide", "duration": "23 hrs", "level": "Intermediate"},
            {"name": "Docker Deep Dive", "platform": "Pluralsight", "url": "https://pluralsight.com", "duration": "6 hrs", "level": "Beginner"},
        ],
        "projects": [
            "Containerize a Flask/Django web application",
            "Set up a multi-container app with Docker Compose",
            "Deploy a microservices architecture with Docker Swarm",
            "Build a CI/CD pipeline with Docker and GitHub Actions",
        ],
        "certifications": [
            {"name": "Docker Certified Associate (DCA)", "org": "Docker Inc.", "url": "https://training.mirantis.com/dca"},
        ],
        "timeline": "2–4 weeks for basics, 2–3 months for production-level"
    },
    "kubernetes": {
        "courses": [
            {"name": "Kubernetes for the Absolute Beginners", "platform": "Udemy (KodeKloud)", "url": "https://udemy.com/course/learn-kubernetes", "duration": "8 hrs", "level": "Beginner"},
            {"name": "Certified Kubernetes Administrator (CKA) Prep", "platform": "KodeKloud", "url": "https://kodekloud.com", "duration": "15 hrs", "level": "Advanced"},
        ],
        "projects": [
            "Deploy a full-stack app on a local Minikube cluster",
            "Set up auto-scaling with HPA on AKS/GKE/EKS",
            "Implement blue-green deployments with Kubernetes",
            "Build a GitOps pipeline with ArgoCD and Kubernetes",
        ],
        "certifications": [
            {"name": "Certified Kubernetes Administrator (CKA)", "org": "CNCF / Linux Foundation", "url": "https://training.linuxfoundation.org"},
            {"name": "Certified Kubernetes Application Developer (CKAD)", "org": "CNCF", "url": "https://training.linuxfoundation.org"},
        ],
        "timeline": "2–4 months of focused practice"
    },
    "tensorflow": {
        "courses": [
            {"name": "TensorFlow: Data and Deployment Specialization", "platform": "Coursera", "url": "https://coursera.org/specializations/tensorflow-data-and-deployment", "duration": "4 months", "level": "Intermediate"},
            {"name": "TensorFlow 2.0 Complete Course", "platform": "freeCodeCamp (YouTube)", "url": "https://youtube.com/watch?v=tPYj3fFJGjk", "duration": "7 hrs", "level": "Beginner"},
        ],
        "projects": [
            "Image classification with CNN using TF/Keras",
            "Time series forecasting with LSTM",
            "Serve a TF model using TensorFlow Serving + Docker",
            "Build a chatbot using TF + Seq2Seq architecture",
        ],
        "certifications": [
            {"name": "TensorFlow Developer Certificate", "org": "Google", "url": "https://tensorflow.org/certificate"},
        ],
        "timeline": "2–4 months (after learning Python & ML)"
    },
    "pytorch": {
        "courses": [
            {"name": "Deep Learning with PyTorch: Zero to GANs", "platform": "Jovian.ai (Free)", "url": "https://jovian.com/learn/deep-learning-with-pytorch-zero-to-gans", "duration": "6 weeks", "level": "Intermediate"},
            {"name": "PyTorch for Deep Learning Bootcamp", "platform": "Udemy", "url": "https://udemy.com/course/pytorch-for-deep-learning-and-computer-vision", "duration": "17 hrs", "level": "Intermediate"},
        ],
        "projects": [
            "Fine-tune a BERT model for text classification",
            "Train a custom ResNet for image recognition",
            "Build a neural style transfer application",
            "Implement a custom training loop with distributed training",
        ],
        "certifications": [
            {"name": "PyTorch Fundamentals Learning Path", "org": "Microsoft Learn (Free)", "url": "https://learn.microsoft.com"},
        ],
        "timeline": "2–4 months (after Python & ML basics)"
    },
    "aws": {
        "courses": [
            {"name": "AWS Cloud Practitioner Essentials", "platform": "AWS Training (Free)", "url": "https://aws.amazon.com/training", "duration": "6 hrs", "level": "Beginner"},
            {"name": "Ultimate AWS Certified Solutions Architect Associate", "platform": "Udemy", "url": "https://udemy.com/course/aws-certified-solutions-architect-associate-saa-c03", "duration": "27 hrs", "level": "Intermediate"},
        ],
        "projects": [
            "Deploy a scalable web app using EC2, ALB, and Auto Scaling",
            "Build a serverless API with AWS Lambda and API Gateway",
            "Set up a data lake with S3, Glue, and Athena",
            "Create a real-time streaming pipeline with Kinesis and Lambda",
        ],
        "certifications": [
            {"name": "AWS Certified Cloud Practitioner", "org": "Amazon Web Services", "url": "https://aws.amazon.com/certification"},
            {"name": "AWS Certified Solutions Architect – Associate", "org": "Amazon Web Services", "url": "https://aws.amazon.com/certification"},
            {"name": "AWS Certified ML – Specialty", "org": "Amazon Web Services", "url": "https://aws.amazon.com/certification"},
        ],
        "timeline": "1–3 months per certification level"
    },
    "react": {
        "courses": [
            {"name": "React – The Complete Guide 2024", "platform": "Udemy", "url": "https://udemy.com/course/react-the-complete-guide-incl-redux", "duration": "68 hrs", "level": "Beginner-Advanced"},
            {"name": "Full Stack Open (React + Node)", "platform": "University of Helsinki (Free)", "url": "https://fullstackopen.com", "duration": "Self-paced", "level": "Intermediate"},
        ],
        "projects": [
            "Build a full-featured e-commerce app with React + Redux",
            "Create a real-time chat application with React + WebSockets",
            "Develop a portfolio website with React + Next.js",
            "Build a task management app with React Query and Zustand",
        ],
        "certifications": [
            {"name": "Meta Front-End Developer Certificate", "org": "Meta / Coursera", "url": "https://coursera.org/professional-certificates/meta-front-end-developer"},
        ],
        "timeline": "2–4 months for professional-level skills"
    },
    "data visualization": {
        "courses": [
            {"name": "Data Visualization with Python", "platform": "Coursera (IBM)", "url": "https://coursera.org/learn/python-for-data-visualization", "duration": "5 weeks", "level": "Beginner"},
            {"name": "Tableau 2024 A-Z", "platform": "Udemy", "url": "https://udemy.com/course/tableau10", "duration": "9 hrs", "level": "Beginner"},
        ],
        "projects": [
            "Create an interactive COVID-19 dashboard with Plotly Dash",
            "Build a sales analytics report in Tableau with filters",
            "Develop a geospatial visualization of Indian census data",
            "Design a real-time stock market dashboard with D3.js",
        ],
        "certifications": [
            {"name": "Tableau Desktop Specialist", "org": "Tableau (Salesforce)", "url": "https://tableau.com/learn/certification"},
            {"name": "Microsoft Power BI Data Analyst (PL-300)", "org": "Microsoft", "url": "https://learn.microsoft.com"},
        ],
        "timeline": "4–8 weeks per tool"
    },
    "nlp": {
        "courses": [
            {"name": "Natural Language Processing Specialization", "platform": "Coursera (deeplearning.ai)", "url": "https://coursera.org/specializations/natural-language-processing", "duration": "4 months", "level": "Intermediate"},
            {"name": "HuggingFace NLP Course", "platform": "HuggingFace (Free)", "url": "https://huggingface.co/learn/nlp-course", "duration": "Self-paced", "level": "Intermediate"},
        ],
        "projects": [
            "Build a resume parser using spaCy NER",
            "Create a customer support chatbot with Rasa",
            "Develop a news article summarizer with BART/T5",
            "Build a semantic search engine with FAISS + Sentence-BERT",
        ],
        "certifications": [
            {"name": "NLP with Classification and Vector Spaces", "org": "deeplearning.ai", "url": "https://coursera.org"},
        ],
        "timeline": "3–5 months (after Python & ML basics)"
    },
    "statistics": {
        "courses": [
            {"name": "Statistics with Python Specialization", "platform": "Coursera (UMich)", "url": "https://coursera.org/specializations/statistics-with-python", "duration": "5 months", "level": "Beginner"},
            {"name": "Statistics and Probability – Khan Academy", "platform": "Khan Academy (Free)", "url": "https://khanacademy.org/math/statistics-probability", "duration": "Self-paced", "level": "Beginner"},
        ],
        "projects": [
            "Conduct an A/B test analysis on an e-commerce dataset",
            "Build a Bayesian inference model from scratch",
            "Perform survival analysis on clinical trial data",
            "Statistical analysis of IPL cricket data with hypothesis tests",
        ],
        "certifications": [
            {"name": "Statistical Learning (free audit)", "org": "Stanford Online", "url": "https://online.stanford.edu"},
        ],
        "timeline": "2–3 months for solid fundamentals"
    },
    "git": {
        "courses": [
            {"name": "Git & GitHub – The Complete Guide", "platform": "Udemy", "url": "https://udemy.com/course/git-github-practical-guide", "duration": "6 hrs", "level": "Beginner"},
            {"name": "Pro Git Book", "platform": "Git-SCM (Free)", "url": "https://git-scm.com/book", "duration": "Self-paced", "level": "All"},
        ],
        "projects": [
            "Contribute to an open-source project on GitHub",
            "Set up GitFlow branching strategy for a team project",
            "Build a CI/CD pipeline with GitHub Actions",
            "Manage a monorepo with Git submodules",
        ],
        "certifications": [
            {"name": "GitHub Foundations Certification", "org": "GitHub", "url": "https://education.github.com"},
        ],
        "timeline": "1–2 weeks for basic proficiency"
    },
    "spark": {
        "courses": [
            {"name": "Apache Spark with Python – PySpark", "platform": "Udemy", "url": "https://udemy.com/course/apache-spark-with-scala-hands-on-with-big-data", "duration": "14 hrs", "level": "Intermediate"},
            {"name": "Databricks Fundamentals", "platform": "Databricks Academy (Free)", "url": "https://academy.databricks.com", "duration": "Self-paced", "level": "Beginner"},
        ],
        "projects": [
            "Build a streaming data pipeline with Spark Structured Streaming",
            "Analyze a 10GB e-commerce dataset using PySpark on Databricks",
            "Implement ML pipelines with Spark MLlib",
        ],
        "certifications": [
            {"name": "Databricks Certified Associate Developer for Apache Spark", "org": "Databricks", "url": "https://databricks.com/certification"},
        ],
        "timeline": "2–3 months for hands-on proficiency"
    },
    "default": {
        "courses": [
            {"name": "Search on Coursera / Udemy / edX", "platform": "Multiple Platforms", "url": "https://coursera.org", "duration": "Varies", "level": "All"},
            {"name": "Official Documentation & Tutorials", "platform": "Official Docs", "url": "#", "duration": "Self-paced", "level": "All"},
            {"name": "YouTube tutorials (freeCodeCamp, Traversy Media, etc.)", "platform": "YouTube (Free)", "url": "https://youtube.com", "duration": "Self-paced", "level": "Beginner"},
        ],
        "projects": [
            "Build a beginner project using this skill",
            "Contribute to open source projects using this skill",
            "Create a portfolio project showcasing this skill",
        ],
        "certifications": [
            {"name": "Look for role-specific certifications on Credly / LinkedIn Learning", "org": "Multiple", "url": "https://credly.com"},
        ],
        "timeline": "1–3 months depending on complexity"
    }
}

# ═══════════════════════════════════════════════════════════════════════════════
#  UTILITY FUNCTIONS
# ═══════════════════════════════════════════════════════════════════════════════

def extract_text_from_pdf(filepath):
    text = ""
    if not PDF_OK:
        return "pdfplumber not installed. Please run: pip install pdfplumber"
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        text = f"Error reading PDF: {str(e)}"
    return text

def extract_text_from_docx(filepath):
    text = ""
    if not DOCX_OK:
        return "python-docx not installed. Please run: pip install python-docx"
    try:
        doc = DocxDocument(filepath)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        text = f"Error reading DOCX: {str(e)}"
    return text

def clean_text(text):
    text = re.sub(r'[^\w\s\.\,\-\+\#]', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text.lower().strip()

def extract_skills_from_text(text):
    """Extract skills by matching against the full job skills database."""
    text_lower = text.lower()
    found = set()
    all_skills = set()
    for job_data in JOB_DATABASE.values():
        all_skills.update(job_data["skills"])

    for skill in all_skills:
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text_lower):
            found.add(skill)
    return sorted(found)

def compute_ats_score(text):
    """Compute ATS score based on presence of key resume sections and keywords."""
    text_lower = text.lower()
    score = 0
    breakdown = {}
    max_scores = {
        "Contact Information": 10,
        "Education": 15,
        "Experience": 20,
        "Technical Skills": 20,
        "Soft Skills": 10,
        "Certifications": 10,
        "Achievements": 10,
        "Action Verbs": 5,
    }
    for section, keywords in ATS_SECTIONS.items():
        found = sum(1 for kw in keywords if kw.lower() in text_lower)
        ratio = min(found / max(len(keywords) * 0.4, 1), 1.0)
        section_score = round(ratio * max_scores[section])
        score += section_score
        breakdown[section] = {
            "score": section_score,
            "max": max_scores[section],
            "found": found,
            "total": len(keywords)
        }

    # Bonus: resume length check
    word_count = len(text.split())
    if 300 <= word_count <= 800:
        score = min(score + 5, 100)
    elif word_count > 800:
        score = max(score - 3, 0)

    # Quantifiable achievements
    if re.search(r'\d+\s*%|\d+\s*x\s|\$\s*\d+|₹\s*\d+|\d+\s*(users|customers|employees|teams)', text_lower):
        score = min(score + 5, 100)

    return min(score, 100), breakdown

def tfidf_cosine_similarity(text1, text2):
    """Compute cosine similarity using TF-IDF (test_size=0.3 equivalent in evaluation)."""
    if not ML_OK:
        # Fallback: Jaccard similarity
        set1 = set(text1.lower().split())
        set2 = set(text2.lower().split())
        intersection = set1 & set2
        union = set1 | set2
        return len(intersection) / len(union) if union else 0.0

    vectorizer = TfidfVectorizer(
        ngram_range=(1, 2),
        stop_words='english',
        min_df=1,
        max_features=5000
    )
    try:
        tfidf_matrix = vectorizer.fit_transform([text1, text2])
        sim = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
        return float(sim)
    except Exception:
        return 0.0

def match_jobs(resume_skills, resume_text):
    """Match resume against all jobs using TF-IDF cosine similarity."""
    results = []
    resume_skills_text = " ".join(resume_skills) + " " + resume_text[:2000]

    for job_title, job_data in JOB_DATABASE.items():
        job_skills_text = " ".join(job_data["skills"])

        # TF-IDF similarity
        similarity = tfidf_cosine_similarity(resume_skills_text, job_skills_text)

        # Skill overlap
        resume_set = set(s.lower() for s in resume_skills)
        job_set = set(s.lower() for s in job_data["skills"])
        matched = resume_set & job_set
        missing = job_set - resume_set
        match_pct = round(len(matched) / len(job_set) * 100, 1) if job_set else 0

        results.append({
            "title": job_title,
            "similarity_score": round(similarity * 100, 1),
            "skill_match_pct": match_pct,
            "matched_skills": sorted(matched),
            "missing_skills": sorted(missing),
            "avg_salary": job_data["avg_salary"],
            "demand": job_data["demand"],
            "description": job_data["description"],
            "companies": job_data["companies"],
            "total_skills_required": len(job_set)
        })

    # Sort by composite score
    results.sort(key=lambda x: (x["similarity_score"] * 0.5 + x["skill_match_pct"] * 0.5), reverse=True)
    return results

def get_skill_gap(resume_skills, target_job):
    """Compute missing skills for a specific job."""
    job_data = JOB_DATABASE.get(target_job)
    if not job_data:
        # Fuzzy match
        target_lower = target_job.lower()
        for title in JOB_DATABASE:
            if target_lower in title.lower() or title.lower() in target_lower:
                job_data = JOB_DATABASE[title]
                target_job = title
                break
    if not job_data:
        return None, None

    resume_set = set(s.lower() for s in resume_skills)
    job_set = set(s.lower() for s in job_data["skills"])
    matched = resume_set & job_set
    missing = job_set - resume_set

    sim = tfidf_cosine_similarity(" ".join(resume_skills), " ".join(job_data["skills"]))

    return {
        "job_title": target_job,
        "similarity": round(sim * 100, 1),
        "matched_skills": sorted(matched),
        "missing_skills": sorted(missing),
        "match_pct": round(len(matched) / len(job_set) * 100, 1) if job_set else 0,
        "total_required": len(job_set),
        "avg_salary": job_data["avg_salary"],
        "demand": job_data["demand"],
        "description": job_data["description"],
        "companies": job_data["companies"],
    }, job_data

def get_learning_roadmap(missing_skills):
    """Map missing skills to curated learning resources."""
    roadmap = []
    for skill in missing_skills:
        skill_lower = skill.lower()
        resources = LEARNING_RESOURCES.get(skill_lower, LEARNING_RESOURCES["default"])
        roadmap.append({
            "skill": skill,
            "courses": resources["courses"],
            "projects": resources["projects"],
            "certifications": resources["certifications"],
            "timeline": resources.get("timeline", "1–3 months")
        })
    return roadmap

def extract_name(text):
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    for line in lines[:5]:
        words = line.split()
        if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w):
            if not any(kw in line.lower() for kw in ['resume', 'cv', 'curriculum', 'vitae', 'summary']):
                return line
    return "Candidate"

def extract_email(text):
    match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    return match.group(0) if match else "N/A"

def extract_phone(text):
    match = re.search(r'(\+91[\s\-]?)?[6-9]\d{9}|(\+1[\s\-]?)?\d{3}[\s\-]\d{3}[\s\-]\d{4}', text)
    return match.group(0) if match else "N/A"

def extract_experience_years(text):
    patterns = [
        r'(\d+)\+?\s*years?\s*of\s*(experience|exp)',
        r'(\d+)\+?\s*yrs?\s*of\s*(experience|exp)',
        r'experience\s*:?\s*(\d+)\+?\s*years?',
    ]
    for p in patterns:
        m = re.search(p, text, re.IGNORECASE)
        if m:
            return int(m.group(1))
    # Count from dates
    years = re.findall(r'\b(20\d{2})\b', text)
    if len(years) >= 2:
        years = sorted(set(int(y) for y in years))
        return max(1, years[-1] - years[0])
    return 0

def skill_categories(skills):
    cats = {
        "Programming Languages": ["python","java","javascript","typescript","c++","c","golang","scala","r","ruby","php","swift","kotlin"],
        "ML/AI Frameworks": ["tensorflow","pytorch","keras","scikit-learn","xgboost","lightgbm","huggingface","transformers","spacy","nltk"],
        "Data & Analytics": ["pandas","numpy","sql","spark","hadoop","tableau","power bi","matplotlib","seaborn","plotly","r","excel"],
        "Cloud & DevOps": ["aws","gcp","azure","docker","kubernetes","terraform","jenkins","ci/cd","git","linux","bash","ansible"],
        "Databases": ["mysql","postgresql","mongodb","redis","cassandra","elasticsearch","bigquery","snowflake","sqlite"],
        "Web & APIs": ["react","nodejs","flask","fastapi","django","html","css","rest api","graphql","nginx"],
        "Soft Skills": ["communication","leadership","teamwork","problem solving","agile","scrum"],
    }
    result = {}
    for cat, cat_skills in cats.items():
        matched = [s for s in skills if s.lower() in cat_skills]
        if matched:
            result[cat] = matched
    uncategorized = [s for s in skills if not any(s.lower() in cs for cs in cats.values())]
    if uncategorized:
        result["Other Skills"] = uncategorized
    return result

# ═══════════════════════════════════════════════════════════════════════════════
#  PDF REPORT GENERATION
# ═══════════════════════════════════════════════════════════════════════════════

def generate_analysis_pdf(data):
    """Generate comprehensive PDF for full analysis report."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            rightMargin=2*cm, leftMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story = []

    C_PRIMARY = colors.HexColor("#00D4FF")
    C_SECONDARY = colors.HexColor("#7B5CFF")
    C_DARK = colors.HexColor("#050A14")
    C_TEXT = colors.HexColor("#1a1a2e")
    C_ACCENT = colors.HexColor("#00FFB3")
    C_WARN = colors.HexColor("#FF6B6B")
    C_GOLD = colors.HexColor("#FFD166")

    title_style = ParagraphStyle('Title', fontName='Helvetica-Bold', fontSize=22,
                                  textColor=C_SECONDARY, spaceAfter=4, alignment=TA_CENTER)
    subtitle_style = ParagraphStyle('Subtitle', fontName='Helvetica', fontSize=11,
                                     textColor=colors.HexColor("#666666"), spaceAfter=20, alignment=TA_CENTER)
    h1_style = ParagraphStyle('H1', fontName='Helvetica-Bold', fontSize=15,
                               textColor=C_SECONDARY, spaceBefore=16, spaceAfter=6)
    h2_style = ParagraphStyle('H2', fontName='Helvetica-Bold', fontSize=12,
                               textColor=C_PRIMARY, spaceBefore=10, spaceAfter=4)
    normal_style = ParagraphStyle('Normal2', fontName='Helvetica', fontSize=9.5,
                                   textColor=C_TEXT, spaceAfter=4, leading=14)
    bullet_style = ParagraphStyle('Bullet', fontName='Helvetica', fontSize=9,
                                   textColor=C_TEXT, spaceAfter=3, leftIndent=12, leading=13)

    # Header
    story.append(Paragraph("ResumeIQ", title_style))
    story.append(Paragraph("AI-Powered Resume Analysis Report", subtitle_style))
    story.append(HRFlowable(width="100%", thickness=2, color=C_PRIMARY, spaceAfter=10))

    meta_data = [
        ["Candidate", data.get("name","N/A"), "Email", data.get("email","N/A")],
        ["Phone", data.get("phone","N/A"), "Experience", f"{data.get('experience_years',0)} yrs"],
        ["Report Date", datetime.now().strftime("%d %b %Y"), "ATS Score", f"{data.get('ats_score',0)}/100"],
    ]
    meta_table = Table(meta_data, colWidths=[3.5*cm, 6.5*cm, 3.5*cm, 4*cm])
    meta_table.setStyle(TableStyle([
        ('FONTNAME', (0,0), (-1,-1), 'Helvetica'),
        ('FONTSIZE', (0,0), (-1,-1), 9),
        ('FONTNAME', (0,0), (0,-1), 'Helvetica-Bold'),
        ('FONTNAME', (2,0), (2,-1), 'Helvetica-Bold'),
        ('TEXTCOLOR', (0,0), (0,-1), C_SECONDARY),
        ('TEXTCOLOR', (2,0), (2,-1), C_SECONDARY),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#dddddd")),
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor("#f7f9ff")),
        ('ROWBACKGROUNDS', (0,0), (-1,-1), [colors.HexColor("#f7f9ff"), colors.white]),
        ('PADDING', (0,0), (-1,-1), 6),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 12))

    # ATS Score
    story.append(Paragraph("ATS Score Breakdown", h1_style))
    ats_score = data.get("ats_score", 0)
    ats_label = "Excellent ✓" if ats_score >= 80 else "Good" if ats_score >= 60 else "Needs Improvement"
    ats_color = C_ACCENT if ats_score >= 80 else C_GOLD if ats_score >= 60 else C_WARN
    story.append(Paragraph(f"Overall ATS Score: <b>{ats_score}/100</b> — {ats_label}", normal_style))
    story.append(Spacer(1, 6))

    ats_bd = data.get("ats_breakdown", {})
    if ats_bd:
        ats_rows = [["Section", "Score", "Max", "Keyword Coverage"]]
        for section, info in ats_bd.items():
            pct = round(info['score'] / info['max'] * 100) if info['max'] else 0
            bar = "█" * (pct // 10) + "░" * (10 - pct // 10)
            ats_rows.append([section, str(info['score']), str(info['max']), f"{bar} {pct}%"])
        ats_table = Table(ats_rows, colWidths=[5.5*cm, 2*cm, 2*cm, 8*cm])
        ats_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), C_SECONDARY),
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,-1), 8.5),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor("#f5f8ff")]),
            ('GRID', (0,0), (-1,-1), 0.4, colors.HexColor("#dddddd")),
            ('PADDING', (0,0), (-1,-1), 5),
        ]))
        story.append(ats_table)
    story.append(Spacer(1, 12))

    # Skills
    story.append(Paragraph("Extracted Skills", h1_style))
    cats = data.get("skill_categories", {})
    for cat, skills in cats.items():
        story.append(Paragraph(cat, h2_style))
        skills_text = " • ".join(skills)
        story.append(Paragraph(skills_text, bullet_style))

    # Top Job Matches
    story.append(Paragraph("Top Job Matches (AI Ranked)", h1_style))
    top_jobs = data.get("top_jobs", [])[:8]
    if top_jobs:
        job_rows = [["Job Title", "Match %", "Skills Matched", "Salary", "Demand"]]
        for j in top_jobs:
            job_rows.append([
                j["title"],
                f"{j['skill_match_pct']}%",
                f"{len(j['matched_skills'])}/{j['total_skills_required']}",
                j["avg_salary"],
                j["demand"]
            ])
        job_table = Table(job_rows, colWidths=[5*cm, 2.5*cm, 3.5*cm, 3.5*cm, 3*cm])
        job_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), C_PRIMARY),
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,-1), 8.5),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor("#f0faff")]),
            ('GRID', (0,0), (-1,-1), 0.4, colors.HexColor("#dddddd")),
            ('PADDING', (0,0), (-1,-1), 5),
        ]))
        story.append(job_table)

    story.append(Spacer(1, 8))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#dddddd")))
    story.append(Paragraph(f"Generated by ResumeIQ • {datetime.now().strftime('%d %b %Y %H:%M')}",
                             ParagraphStyle('Footer', fontName='Helvetica', fontSize=8,
                                            textColor=colors.HexColor("#999999"), alignment=TA_CENTER)))
    doc.build(story)
    buffer.seek(0)
    return buffer

def generate_gap_pdf(data):
    """Generate PDF for skill gap analysis."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            rightMargin=2*cm, leftMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story = []

    C_PRIMARY = colors.HexColor("#00D4FF")
    C_SECONDARY = colors.HexColor("#7B5CFF")
    C_WARN = colors.HexColor("#FF6B6B")
    C_ACCENT = colors.HexColor("#00FFB3")
    C_GOLD = colors.HexColor("#FFD166")

    title_style = ParagraphStyle('Title', fontName='Helvetica-Bold', fontSize=22,
                                  textColor=C_SECONDARY, spaceAfter=4, alignment=TA_CENTER)
    subtitle_style = ParagraphStyle('Subtitle', fontName='Helvetica', fontSize=11,
                                     textColor=colors.HexColor("#666666"), spaceAfter=20, alignment=TA_CENTER)
    h1_style = ParagraphStyle('H1', fontName='Helvetica-Bold', fontSize=14,
                               textColor=C_SECONDARY, spaceBefore=14, spaceAfter=6)
    h2_style = ParagraphStyle('H2', fontName='Helvetica-Bold', fontSize=11,
                               textColor=C_PRIMARY, spaceBefore=8, spaceAfter=4)
    h3_style = ParagraphStyle('H3', fontName='Helvetica-Bold', fontSize=10,
                               textColor=colors.HexColor("#333333"), spaceBefore=6, spaceAfter=3)
    normal_style = ParagraphStyle('Normal2', fontName='Helvetica', fontSize=9.5,
                                   textColor=colors.HexColor("#1a1a2e"), spaceAfter=4, leading=14)
    bullet_style = ParagraphStyle('Bullet', fontName='Helvetica', fontSize=9,
                                   textColor=colors.HexColor("#1a1a2e"), spaceAfter=2,
                                   leftIndent=15, leading=13)

    story.append(Paragraph("ResumeIQ", title_style))
    story.append(Paragraph("Skill Gap Analysis & Learning Roadmap", subtitle_style))
    story.append(HRFlowable(width="100%", thickness=2, color=C_PRIMARY, spaceAfter=12))

    gap = data.get("gap", {})
    story.append(Paragraph(f"Target Role: {gap.get('job_title','N/A')}", h1_style))

    summary_data = [
        ["Metric", "Value"],
        ["Skills Matched", f"{len(gap.get('matched_skills',[]))} / {gap.get('total_required',0)}"],
        ["Match Percentage", f"{gap.get('match_pct',0)}%"],
        ["Skills to Acquire", str(len(gap.get('missing_skills',[])))],
        ["AI Similarity Score", f"{gap.get('similarity',0)}%"],
        ["Average Salary", gap.get("avg_salary","N/A")],
        ["Market Demand", gap.get("demand","N/A")],
    ]
    summary_table = Table(summary_data, colWidths=[8*cm, 9.5*cm])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), C_SECONDARY),
        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 9.5),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor("#f5f8ff")]),
        ('GRID', (0,0), (-1,-1), 0.4, colors.HexColor("#cccccc")),
        ('PADDING', (0,0), (-1,-1), 7),
        ('FONTNAME', (0,1), (0,-1), 'Helvetica-Bold'),
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 12))

    # Skills you have
    story.append(Paragraph("✅ Skills You Already Have", h1_style))
    matched = gap.get("matched_skills", [])
    if matched:
        story.append(Paragraph(" • ".join(matched), bullet_style))
    else:
        story.append(Paragraph("No matching skills found.", normal_style))
    story.append(Spacer(1, 8))

    # Missing skills
    story.append(Paragraph("❌ Skills You Need to Acquire", h1_style))
    missing = gap.get("missing_skills", [])
    if missing:
        story.append(Paragraph(" • ".join(missing), bullet_style))
    else:
        story.append(Paragraph("No missing skills — you are fully qualified!", normal_style))
    story.append(Spacer(1, 12))

    # Learning roadmap
    story.append(Paragraph("📚 Personalised Learning Roadmap", h1_style))
    roadmap = data.get("roadmap", [])
    for i, item in enumerate(roadmap, 1):
        story.append(Paragraph(f"{i}. {item['skill'].upper()}", h2_style))
        story.append(Paragraph(f"Estimated Timeline: {item['timeline']}", normal_style))

        story.append(Paragraph("Recommended Courses:", h3_style))
        for c in item["courses"]:
            story.append(Paragraph(f"• {c['name']} ({c['platform']}) — {c['duration']} [{c['level']}]",
                                    bullet_style))

        story.append(Paragraph("Hands-on Projects:", h3_style))
        for p in item["projects"][:3]:
            story.append(Paragraph(f"• {p}", bullet_style))

        story.append(Paragraph("Certifications:", h3_style))
        for cert in item["certifications"]:
            story.append(Paragraph(f"• {cert['name']} — {cert['org']}", bullet_style))
        story.append(Spacer(1, 6))

    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#dddddd")))
    story.append(Paragraph(f"Generated by ResumeIQ • {datetime.now().strftime('%d %b %Y %H:%M')}",
                             ParagraphStyle('Footer', fontName='Helvetica', fontSize=8,
                                            textColor=colors.HexColor("#999999"), alignment=TA_CENTER)))
    doc.build(story)
    buffer.seek(0)
    return buffer

# ═══════════════════════════════════════════════════════════════════════════════
#  ROUTES
# ═══════════════════════════════════════════════════════════════════════════════

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/test_db')
def test_db():
    cur = mysql.connection.cursor()
    cur.execute("SELECT 1")
    return "Database Connected Successfully!"

# ── Auth ──────────────────────────────────────────────────────────────────────
@app.route('/api/login', methods=['POST'])
def login():
    try:
        data = request.get_json()
        email = data.get('email', '').strip()
        password = data.get('password', '').strip()

        if not email or not password:
            return jsonify({"success": False, "message": "Email and password are required."})

        cur = mysql.connection.cursor()

        # ✅ fetch full user
        cur.execute("SELECT id, name, password FROM users WHERE email=%s", (email,))
        user = cur.fetchone()

        cur.close()

        # ❌ user not found
        if not user:
            return jsonify({"success": False, "message": "User not found"})

        # ❌ password mismatch
        if password != user[2]:
            return jsonify({"success": False, "message": "Incorrect password"})

        # ✅ success
        session['user_id'] = user[0]
        session['name'] = user[1]

        return jsonify({"success": True, "name": user[1]})

    except Exception as e:
        print("LOGIN ERROR:", str(e))
        return jsonify({"success": False, "message": str(e)})

@app.route('/api/register', methods=['POST'])
def register():
    try:
        data = request.get_json()
        email = data.get('email', '').strip()
        password = data.get('password', '').strip()
        name = data.get('name', '').strip()

        if not all([email, password, name]):
            return jsonify({"success": False, "message": "All fields are required."})

        if '@' not in email:
            return jsonify({"success": False, "message": "Invalid email format."})

        if len(password) < 6:
            return jsonify({"success": False, "message": "Password must be at least 6 characters."})

        cur = mysql.connection.cursor()

        # ✅ check if user already exists
        cur.execute("SELECT id FROM users WHERE email=%s", (email,))
        existing = cur.fetchone()

        if existing:
            cur.close()
            return jsonify({"success": False, "message": "User already exists"})

        # ✅ insert user into DB
        cur.execute(
            "INSERT INTO users (name, email, password) VALUES (%s, %s, %s)",
            (name, email, password)
        )
        mysql.connection.commit()

        user_id = cur.lastrowid
        cur.close()

        # ✅ store session
        session['user_id'] = user_id
        session['name'] = name

        return jsonify({"success": True, "name": name})

    except Exception as e:
        print("REGISTER ERROR:", str(e))
        return jsonify({"success": False, "message": str(e)})

@app.route('/api/logout', methods=['POST'])
def logout():
    session.clear()
    return jsonify({"success": True})

# ── Resume Upload ─────────────────────────────────────────────────────────────
@app.route('/api/upload', methods=['POST'])
def upload_resume():
    
    if 'user_id' not in session:
        return jsonify({"success": False, "message": "Please login first."})
    if 'resume' not in request.files:
        return jsonify({"success": False, "message": "No file uploaded."})
    file = request.files['resume']
    if file.filename == '':
        return jsonify({"success": False, "message": "No file selected."})

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ['.pdf', '.docx']:
        return jsonify({"success": False, "message": "Only PDF and DOCX files are supported."})

    filename = f"{uuid.uuid4().hex}{ext}"
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)

    if ext == '.pdf':
        text = extract_text_from_pdf(filepath)
    else:
        text = extract_text_from_docx(filepath)

    if not text or len(text.strip()) < 50:
        return jsonify({"success": False, "message": "Could not extract text from the file. Please ensure it's not scanned."})
    
    cur = mysql.connection.cursor()

    cur.execute("""
        INSERT INTO resumes (user_id, file_name, extracted_text)
        VALUES (%s, %s, %s)
    """, (session['user_id'], filename, text))

    mysql.connection.commit()
    cur.close()

    session['resume_text'] = text
    session['resume_file'] = filename
    session['resume_name'] = file.filename

    return jsonify({
        "success": True,
        "filename": file.filename,
        "word_count": len(text.split()),
        "char_count": len(text)
    })

# ── Full Analysis ─────────────────────────────────────────────────────────────
@app.route('/api/analyze', methods=['POST'])
def analyze():
    text = session.get('resume_text','')
    if not text:
        return jsonify({"success": False, "message": "Please upload a resume first."})

    name = extract_name(text)
    email = extract_email(text)
    phone = extract_phone(text)
    exp_years = extract_experience_years(text)
    skills = extract_skills_from_text(text)
    ats_score, ats_breakdown = compute_ats_score(text)
    skill_cats = skill_categories(skills)
    job_matches = match_jobs(skills, text)
    top_jobs = job_matches[:10]

    # Skill frequency for chart (top 15 skills by number of jobs they appear in)
    skill_freq = Counter()
    for job in JOB_DATABASE.values():
        for sk in skills:
            if sk in job["skills"]:
                skill_freq[sk] += 1
    top_skills_chart = [{"skill": k, "count": v} for k, v in skill_freq.most_common(15)]

    result = {
        "success": True,
        "name": name,
        "email": email,
        "phone": phone,
        "experience_years": exp_years,
        "total_skills": len(skills),
        "skills": skills,
        "skill_categories": skill_cats,
        "ats_score": ats_score,
        "ats_breakdown": ats_breakdown,
        "top_jobs": top_jobs,
        "top_skills_chart": top_skills_chart,
        "job_demand_distribution": {
            "Very High": sum(1 for j in top_jobs if j["demand"] == "Very High"),
            "High": sum(1 for j in top_jobs if j["demand"] == "High"),
            "Medium": sum(1 for j in top_jobs if j["demand"] == "Medium"),
        }
    }
    session['last_analysis'] = json.dumps(result)
    
    try:
        cur = mysql.connection.cursor()

        # get latest resume id safely
        cur.execute("SELECT id FROM resumes ORDER BY id DESC LIMIT 1")
        row = cur.fetchone()
        resume_id = row[0] if row else None

        # insert only if resume exists
        if resume_id:
            cur.execute("""
                INSERT INTO analysis (resume_id, skills_found, missing_skills, score)
                VALUES (%s, %s, %s, %s)
            """, (
                resume_id,
                ", ".join(skills),
                ", ".join(top_jobs[0]['missing_skills']) if top_jobs else "",
                ats_score
            ))

            mysql.connection.commit()

        cur.close()

    except Exception as e:
        print("ANALYSIS ERROR:", str(e))
    return jsonify(result)

# ── Skill Gap Analysis ────────────────────────────────────────────────────────
@app.route('/api/skill-gap', methods=['POST'])
def skill_gap():
    text = session.get('resume_text','')
    if not text:
        return jsonify({"success": False, "message": "Please upload a resume first."})

    data = request.get_json()
    target_job = data.get('job','').strip()
    if not target_job:
        return jsonify({"success": False, "message": "Please enter a target job title."})

    skills = extract_skills_from_text(text)
    gap, job_data = get_skill_gap(skills, target_job)
    if not gap:
        available = list(JOB_DATABASE.keys())
        return jsonify({
            "success": False,
            "message": f"Job '{target_job}' not found. Available jobs: {', '.join(available[:10])}...",
            "available_jobs": available
        })

    roadmap = get_learning_roadmap(gap["missing_skills"])
    result = {
        "success": True,
        "gap": gap,
        "roadmap": roadmap,
        "available_jobs": list(JOB_DATABASE.keys())
    }
    session['last_gap'] = json.dumps(result)
    return jsonify(result)

# ── PDF Download Endpoints ────────────────────────────────────────────────────
@app.route('/api/download/analysis', methods=['GET'])
def download_analysis():
    last = session.get('last_analysis')
    if not last:
        return jsonify({"error": "No analysis data. Please run analysis first."}), 400
    data = json.loads(last)
    if not REPORTLAB_OK:
        return jsonify({"error": "reportlab not installed. Run: pip install reportlab"}), 500
    buf = generate_analysis_pdf(data)
    return send_file(buf, mimetype='application/pdf',
                     as_attachment=True,
                     download_name=f"ResumeIQ_Analysis_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf")

@app.route('/api/download/gap', methods=['GET'])
def download_gap():
    last = session.get('last_gap')
    if not last:
        return jsonify({"error": "No gap analysis data. Please run skill gap analysis first."}), 400
    data = json.loads(last)
    if not REPORTLAB_OK:
        return jsonify({"error": "reportlab not installed. Run: pip install reportlab"}), 500
    buf = generate_gap_pdf(data)
    return send_file(buf, mimetype='application/pdf',
                     as_attachment=True,
                     download_name=f"ResumeIQ_SkillGap_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf")

@app.route('/api/jobs', methods=['GET'])
def get_jobs():
    return jsonify({"jobs": list(JOB_DATABASE.keys())})

@app.route('/api/my-resumes', methods=['GET'])
def my_resumes():
    user_id = session.get('user_id')

    if not user_id:
        return jsonify({"success": False, "message": "Not logged in"})

    cur = mysql.connection.cursor()

    cur.execute("""
        SELECT id, file_name, uploaded_at
        FROM resumes
        WHERE user_id = %s
        ORDER BY uploaded_at DESC
    """, (user_id,))

    resumes = cur.fetchall()
    cur.close()

    result = []
    for r in resumes:
        result.append({
            "id": r[0],
            "file_name": r[1],
            "uploaded_at": str(r[2])
        })

    return jsonify({"success": True, "resumes": result})

if __name__ == '__main__':
    app.run(host="0.0.0.0", port=5000)
